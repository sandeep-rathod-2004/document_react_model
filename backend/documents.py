from fastapi import APIRouter, File, UploadFile, Form, Depends, HTTPException
from sqlalchemy.orm import Session
from .database import get_db
from .models import Document
import os, shutil, io

# Text extraction libraries
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
import pandas as pd

router = APIRouter()
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def extract_text(file_path: str):
    ext = file_path.split(".")[-1].lower()
    text = ""

    try:
        if ext == "pdf":
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text() or ""

        elif ext in ["docx", "doc"]:
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"

        elif ext in ["txt", "csv"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()

        elif ext in ["xlsx"]:
            df = pd.read_excel(file_path)
            text = "\n".join(df.astype(str).fillna("").values.flatten().tolist())

        elif ext in ["png", "jpg", "jpeg", "gif", "bmp", "tiff"]:
            # OCR for image files
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)

        else:
            text = f"[Unsupported file type: {ext}]"

    except Exception as e:
        print("Text extraction error:", e)
        text = "[Error extracting text]"

    return text.strip() or "[No readable text found]"

@router.post("/")
async def upload_document(
    file: UploadFile = File(...),
    owner_email: str = Form(...),
    db: Session = Depends(get_db)
):
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)

    # Save uploaded file
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"File save error: {e}")

    # Extract text
    text_content = extract_text(file_path)

    # Save to DB
    new_doc = Document(
        filename=file.filename,
        path=file_path,
        owner_email=owner_email,
        content=text_content
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)

    return {"message": "File uploaded successfully", "document": new_doc}

@router.get("/")
def get_documents(owner_email: str, db: Session = Depends(get_db)):
    docs = db.query(Document).filter(Document.owner_email == owner_email).all()
    return docs
